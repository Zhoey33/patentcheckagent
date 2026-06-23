"""这个文件用于按文档格式分发 PDF 与 Word 文本抽取逻辑。"""

from pathlib import Path

from docx import Document

from app.services.errors import UserFacingError
from app.services.pdf_extractor import extract_pdf_text

SUPPORTED_SUFFIXES = {".pdf", ".docx"}
SUPPORTED_FORMAT_LABEL = "PDF 或 Word（.docx）"


def extract_docx_text(path: Path) -> str:
    """Extract text from a Word .docx file."""

    try:
        document = Document(path)
    except Exception as exc:
        raise UserFacingError("Word 文件文本抽取失败，请确认文件未损坏且为 .docx 格式。") from exc

    parts: list[str] = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise UserFacingError("未能抽取到 Word 文本，请确认文档包含可复制文本。")
    return text


def extract_document_text(path: Path) -> str:
    """Extract text from a supported patent application document."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    raise UserFacingError(f"{path.name} 格式不支持，请上传 {SUPPORTED_FORMAT_LABEL} 文件。")
