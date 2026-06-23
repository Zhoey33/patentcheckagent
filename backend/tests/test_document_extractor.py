"""这个文件用于验证多格式文档文本抽取服务。"""

from pathlib import Path

from docx import Document

from app.services.document_extractor import extract_document_text


def test_extract_document_text_reads_docx_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "claims.docx"
    document = Document()
    document.add_paragraph("一种数据处理方法。")
    document.add_paragraph("包括获取输入数据并生成处理结果。")
    document.save(docx_path)

    text = extract_document_text(docx_path)

    assert "一种数据处理方法。" in text
    assert "生成处理结果" in text
